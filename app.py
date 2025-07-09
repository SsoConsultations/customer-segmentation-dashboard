import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.preprocessing import StandardScaler, OneHotEncoder
from sklearn.cluster import KMeans, DBSCAN, AgglomerativeClustering
from sklearn.mixture import GaussianMixture
from sklearn.metrics import silhouette_score, davies_bouldin_score, calinski_harabasz_score
from sklearn.decomposition import PCA
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
import io
import os
import warnings
from PIL import Image # For loading image

# Suppress warnings for cleaner app output
warnings.filterwarnings("ignore")

# --- Session State Initialization ---
if "analysis_completed" not in st.session_state:
    st.session_state.analysis_completed = False
if "file_uploader_key" not in st.session_state:
    st.session_state.file_uploader_key = 0
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

# Page config
st.set_page_config(
    layout="wide",
    page_title="Unsupervised Learning (using Machine Learning)" # Updated title
)

# --- Authentication Logic ---
def login_page():
    # Attempt to load logo for login page
    try:
        logo = Image.open("SsoLogo.jpg")
        st.image(logo, width=150) # Adjust width as needed
    except FileNotFoundError:
        st.warning("SsoLogo.jpg not found. Please ensure it's in the same directory as the script.")
    
    st.title("Login to Unsupervised Learning App")
    st.markdown("Please enter your credentials to access the application.")

    username = st.text_input("Username")
    password = st.text_input("Password", type="password")

    # Access secrets
    correct_username = os.environ.get("STREAMLIT_USERNAME", st.secrets["streamlit_username"])
    correct_password = os.environ.get("STREAMLIT_PASSWORD", st.secrets["streamlit_password"])

    if st.button("Login"):
        if username == correct_username and password == correct_password:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("Incorrect username or password.")

# --- Utility Functions ---

# Function to safely load data
@st.cache_data(show_spinner=False)
def load_data(uploaded_file):
    with st.spinner("Loading data..."):
        try:
            if uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file)
            elif uploaded_file.name.endswith(('.xlsx', '.xls')):
                df = pd.read_excel(uploaded_file)
            else:
                st.error("Unsupported file type. Please upload a .csv or .xlsx file.")
                return None
            return df
        except Exception as e:
            st.error(f"Error loading file: {e}")
            return None

# Function to preprocess data
@st.cache_data(show_spinner=False)
def preprocess_data(df, selected_numeric_cols, selected_categorical_cols, missing_data_strategy, outlier_handling_method):
    with st.spinner("Preprocessing data..."):
        df_proc = df[selected_numeric_cols + selected_categorical_cols].copy()

        # Handle Missing Data
        if missing_data_strategy == "drop_rows":
            initial_rows = df_proc.shape[0]
            df_proc.dropna(inplace=True)
            if df_proc.shape[0] < initial_rows:
                st.warning(f"Dropped {initial_rows - df_proc.shape[0]} rows due to missing values.")
        elif missing_data_strategy == "mean_imputation" and selected_numeric_cols:
            for col in selected_numeric_cols:
                if df_proc[col].isnull().any():
                    df_proc[col].fillna(df_proc[col].mean(), inplace=True)
        elif missing_data_strategy == "mode_imputation" and selected_categorical_cols:
            for col in selected_categorical_cols:
                if df_proc[col].isnull().any():
                    df_proc[col].fillna(df_proc[col].mode()[0], inplace=True)
        elif missing_data_strategy == "median_imputation" and selected_numeric_cols:
            for col in selected_numeric_cols:
                if df_proc[col].isnull().any():
                    df_proc[col].fillna(df_proc[col].median(), inplace=True)

        # Outlier Handling (Simple capping for numeric, based on IQR)
        if outlier_handling_method == "iqr_capping" and selected_numeric_cols:
            for col in selected_numeric_cols:
                Q1 = df_proc[col].quantile(0.25)
                Q3 = df_proc[col].quantile(0.75)
                IQR = Q3 - Q1
                lower_bound = Q1 - 1.5 * IQR
                upper_bound = Q3 + 1.5 * IQR
                df_proc[col] = np.where(df_proc[col] < lower_bound, lower_bound, df_proc[col])
                df_proc[col] = np.where(df_proc[col] > upper_bound, upper_bound, df_proc[col])
                
        # Store original numeric values for profiling later, before scaling
        df_profile = df_proc.copy()

        # Scale Numeric Features
        scaler = StandardScaler()
        if selected_numeric_cols:
            df_proc[selected_numeric_cols] = scaler.fit_transform(df_proc[selected_numeric_cols])

        # One-Hot Encode Categorical Features
        encoder = OneHotEncoder(handle_unknown='ignore', sparse_output=False)
        if selected_categorical_cols:
            encoded_features = encoder.fit_transform(df_proc[selected_categorical_cols])
            encoded_df = pd.DataFrame(encoded_features, columns=encoder.get_feature_names_out(selected_categorical_cols), index=df_proc.index)
            df_proc = pd.concat([df_proc.drop(columns=selected_categorical_cols), encoded_df], axis=1)
        
        # Ensure the index is preserved for alignment with original data later
        scaled_df = df_proc.copy()

        return scaled_df, df_profile # Return scaled_df for clustering and df_profile for original feature values

# Function to run clustering
@st.cache_data(show_spinner=False)
def run_clustering(scaled_df, n_clusters, chosen_algo, eps=None, min_samples=None):
    with st.spinner(f"Running {chosen_algo} clustering..."):
        labels = np.array([])
        if chosen_algo == "KMeans":
            if n_clusters < 2:
                st.error("K-Means requires n_clusters >= 2.")
                return None
            model = KMeans(n_clusters=n_clusters, random_state=42, n_init='auto')
            labels = model.fit_predict(scaled_df)
        elif chosen_algo == "Gaussian Mixture Model":
            if n_clusters < 1: # GMM can have 1 component
                st.error("Gaussian Mixture Model requires n_components >= 1.")
                return None
            model = GaussianMixture(n_components=n_clusters, random_state=42)
            labels = model.fit_predict(scaled_df)
        elif chosen_algo == "Agglomerative Clustering":
            if n_clusters < 2:
                st.error("Agglomerative Clustering requires n_clusters >= 2.")
                return None
            model = AgglomerativeClustering(n_clusters=n_clusters)
            labels = model.fit_predict(scaled_df)
        elif chosen_algo == "DBSCAN":
            if eps is None or min_samples is None:
                st.error("DBSCAN requires 'eps' and 'min_samples' parameters.")
                return None
            model = DBSCAN(eps=eps, min_samples=min_samples)
            labels = model.fit_predict(scaled_df)
        
        return labels

# Function to evaluate clustering
@st.cache_data(show_spinner=False)
def evaluate_clustering(scaled_df, labels):
    with st.spinner("Evaluating clusters..."):
        silhouette, davies_bouldin, calinski_harabasz = [np.nan] * 3
        
        # Silhouette Score requires at least 2 clusters and > 1 sample
        if len(np.unique(labels)) > 1 and len(labels) > 1:
            try:
                silhouette = silhouette_score(scaled_df, labels)
            except ValueError:
                silhouette = np.nan # In case of single-cluster or other issues
        
        # Davies-Bouldin Index requires at least 2 clusters
        if len(np.unique(labels)) > 1:
            try:
                davies_bouldin = davies_bouldin_score(scaled_df, labels)
            except ValueError:
                davies_bouldin = np.nan
        
        # Calinski-Harabasz Index requires at least 2 clusters and > 1 sample
        if len(np.unique(labels)) > 1 and len(labels) > 1:
            try:
                calinski_harabasz = calinski_harabasz_score(scaled_df, labels)
            except ValueError:
                calinski_harabasz = np.nan

        return silhouette, davies_bouldin, calinski_harabasz

# Function to generate plots and cluster statistics
@st.cache_data(show_spinner=False)
def generate_plots(df, labels, numeric_cols, pca_components=2):
    with st.spinner("Generating visualizations and cluster profiles..."):
        df_clustered = df.copy()
        df_clustered['Cluster'] = labels

        # Ensure numeric_cols used for PCA are actually present in df_clustered
        valid_numeric_cols = [col for col in numeric_cols if col in df_clustered.columns]

        fig_pca = None
        if valid_numeric_cols and len(valid_numeric_cols) >= pca_components:
            try:
                pca = PCA(n_components=pca_components)
                components = pca.fit_transform(df_clustered[valid_numeric_cols])
                pca_df = pd.DataFrame(data=components, columns=[f'PC{i+1}' for i in range(pca_components)])
                pca_df['Cluster'] = df_clustered['Cluster'].values # Aligning by values
                
                fig_pca, ax_pca = plt.subplots(figsize=(10, 7))
                sns.scatterplot(x='PC1', y='PC2', hue='Cluster', data=pca_df, 
                                palette='viridis', legend='full', ax=ax_pca)
                ax_pca.set_title('PCA of Clustered Data')
                plt.close(fig_pca) # Prevent showing plot immediately
            except Exception as e:
                st.warning(f"Could not generate PCA plot: {e}")
        elif len(valid_numeric_cols) < pca_components and len(valid_numeric_cols) > 0:
             st.warning(f"Not enough numeric features ({len(valid_numeric_cols)}) for {pca_components} PCA components. Skipping PCA plot.")
        else:
            st.warning("No valid numeric features selected for PCA plot.")


        fig_profile_numeric = None
        cluster_means_numeric = pd.DataFrame()
        if valid_numeric_cols:
            cluster_means_numeric = df_clustered.groupby('Cluster')[valid_numeric_cols].mean()
            # Normalize for radar chart if needed, or simply plot means
            fig_profile_numeric, axes = plt.subplots(ncols=len(valid_numeric_cols), figsize=(12, 4), sharey=True)
            if len(valid_numeric_cols) == 1: # Handle single subplot case
                axes = [axes]
            for i, col in enumerate(valid_numeric_cols):
                sns.barplot(x=cluster_means_numeric.index, y=cluster_means_numeric[col], ax=axes[i], palette='viridis')
                axes[i].set_title(col)
                axes[i].set_xlabel('Cluster')
                axes[i].set_ylabel('Mean Value' if i == 0 else '')
            plt.tight_layout()
            plt.close(fig_profile_numeric) # Prevent showing plot immediately
        else:
            st.warning("No numeric features to profile.")

        cluster_cat_proportions = pd.DataFrame()
        categorical_cols = [col for col in df.columns if df[col].dtype == 'object' and col != 'Cluster']
        if categorical_cols:
            # Assuming original categorical columns are still present in 'df' (df_profile)
            # and 'df_clustered' contains the 'Cluster' column.
            # We need to ensure we work with the original categorical columns, not one-hot encoded ones.
            # The 'df' passed to this function is 'df_profile_with_labels', which still has original categorical values.
            
            # Select original categorical columns from df_clustered (which is df_profile_with_labels)
            valid_cat_cols_for_profile = [col for col in categorical_cols if col in df_clustered.columns]

            if valid_cat_cols_for_profile:
                # Calculate proportions for each category within each cluster
                proportions_list = []
                for col in valid_cat_cols_for_profile:
                    prop_df = df_clustered.groupby('Cluster')[col].value_counts(normalize=True).unstack(fill_value=0)
                    prop_df = prop_df.apply(lambda x: x / x.sum(), axis=1) # Ensure sums to 1 across categories for each cluster
                    prop_df = prop_df.stack().reset_index(name='Proportion')
                    prop_df['Feature'] = col
                    proportions_list.append(prop_df)
                
                if proportions_list:
                    cluster_cat_proportions = pd.concat(proportions_list, ignore_index=True)
                    # Convert to a more readable pivot format for the report if needed
                    # For a summary table, we might pivot it differently or just display this raw table.
                    # For the report, it's better to show it as proportions for each cluster and category.
                    # Re-pivot for a more suitable table output:
                    pivot_table_data = {}
                    for col in valid_cat_cols_for_profile:
                        # Select only the current feature's data
                        feature_data = cluster_cat_proportions[cluster_cat_proportions['Feature'] == col]
                        # Pivot to get Cluster as index, and categories as columns
                        pivot = feature_data.pivot(index='Cluster', columns='level_1', values='Proportion').fillna(0)
                        # Rename columns to include feature name for clarity
                        pivot.columns = [f"{col}: {cat}" for cat in pivot.columns]
                        pivot_table_data[col] = pivot
                    
                    if pivot_table_data:
                        # Combine all categorical feature pivots into a single DataFrame for easier reporting
                        cluster_cat_proportions = pd.concat(pivot_table_data.values(), axis=1).sort_index()

        return fig_pca, fig_profile_numeric, cluster_means_numeric, cluster_cat_proportions

def format_percentage(value):
    if pd.isna(value):
        return "N/A"
    return f"{value:.1%}"

def format_number(value):
    if pd.isna(value):
        return "N/A"
    return f"{value:.2f}"

def format_value_for_report(value):
    if isinstance(value, (int, float)):
        return f"{value:.2f}"
    return str(value)

# Function to generate cluster summaries (textual descriptions)
def generate_cluster_summaries(cluster_means_numeric, cluster_cat_proportions, df_data_with_labels, labels, chosen_algo):
    structured_summaries = {}
    
    # Calculate overall means and stds for numeric features from the provided df_data_with_labels
    overall_means = None
    overall_stds = None
    if cluster_means_numeric is not None and not cluster_means_numeric.empty:
        # Exclude 'Cluster' column if present in df_data_with_labels
        numeric_cols_for_overall_stats = [col for col in cluster_means_numeric.columns if col in df_data_with_labels.columns and col != 'Cluster']
        if numeric_cols_for_overall_stats:
            overall_means = df_data_with_labels[numeric_cols_for_overall_stats].mean()
            overall_stds = df_data_with_labels[numeric_cols_for_overall_stats].std()
        
    # Get counts for each cluster from the labels array
    unique_labels, counts = np.unique(labels, return_counts=True)
    label_counts = dict(zip(unique_labels, counts))
    
    # Determine the clusters to iterate over, excluding DBSCAN noise if applicable
    clusters_to_summarize = [lbl for lbl in sorted(unique_labels) if lbl != -1] # Exclude -1 for DBSCAN noise

    total_samples = len(labels) # Total samples in the labels array (after preprocessing)

    for cluster_id_int in clusters_to_summarize:
        cluster_id_str = str(cluster_id_int) # For dictionary keys and indexing
        cluster_size = label_counts.get(cluster_id_int, 0) # Get size from the labels array

        if cluster_size == 0: # This handles cases where a cluster might genuinely have 0 members
            structured_summaries[cluster_id_str] = {
                "cluster_heading": f"Cluster {cluster_id_int}: (N={cluster_size}, {0.0:.1f}% of total data)",
                "numeric_characteristics": [],
                "categorical_characteristics": [],
                "persona_implications": "This cluster has 0 members after data processing and therefore no characteristics to display."
            }
            continue

        cluster_percentage = (cluster_size / total_samples) * 100 if total_samples > 0 else 0

        numeric_char = []
        if cluster_means_numeric is not None and cluster_id_str in cluster_means_numeric.index:
            cluster_mean_row = cluster_means_numeric.loc[cluster_id_str]
            for col in cluster_mean_row.index:
                if overall_means is not None and col in overall_means.index and overall_stds is not None and col in overall_stds.index:
                    cluster_mean = cluster_mean_row[col]
                    overall_mean = overall_means[col]
                    overall_std = overall_stds[col]

                    if overall_std > 0: # Avoid division by zero
                        z_score = (cluster_mean - overall_mean) / overall_std
                        # Define thresholds for "significantly higher/lower"
                        # For simplicity, using a +/- 1.0 standard deviation threshold
                        if z_score > 1.0:
                            numeric_char.append(f"- **{col}**: Significantly higher (Avg: {cluster_mean:.2f})")
                        elif z_score < -1.0:
                            numeric_char.append(f"- **{col}**: Significantly lower (Avg: {cluster_mean:.2f})")
                        else:
                            numeric_char.append(f"- **{col}**: Average (Avg: {cluster_mean:.2f})")
                    else: # Overall std is 0, meaning all values are the same
                        numeric_char.append(f"- **{col}**: Consistent (Avg: {cluster_mean:.2f})")
                else:
                    numeric_char.append(f"- **{col}**: Average: {cluster_mean:.2f}")


        categorical_char = []
        # Process categorical features directly from the consolidated proportions DataFrame
        if cluster_cat_proportions is not None and not cluster_cat_proportions.empty:
            # Filter for the current cluster
            cluster_cat_data = cluster_cat_proportions.loc[cluster_id_str]

            # Iterate through each original categorical feature
            # We need to know the original categorical feature names to group them
            # This requires access to the original categorical columns used for encoding
            # For now, let's parse them from the column names of cluster_cat_proportions
            unique_cat_features = sorted(list(set([col.split(': ')[0] for col in cluster_cat_proportions.columns if ': ' in col])))

            for feature_name in unique_cat_features:
                # Find all columns belonging to this feature for the current cluster
                feature_cols = [col for col in cluster_cat_data.index if col.startswith(f"{feature_name}: ")]
                
                if feature_cols:
                    # Get proportions for this cluster for the current categorical feature
                    proportions_for_feature = cluster_cat_data[feature_cols]
                    
                    if not proportions_for_feature.empty:
                        # Find the category with the highest proportion for this cluster
                        most_common_category_full_name = proportions_for_feature.idxmax()
                        most_common_proportion = proportions_for_feature.max()
                        
                        # Extract just the category name from "FeatureName: Category"
                        most_common_category = most_common_category_full_name.split(': ')[1] if ': ' in most_common_category_full_name else most_common_category_full_name
                        
                        # Add to the list if proportion is significant (e.g., > 50%) or if it's the dominant one
                        # You can adjust this threshold
                        if most_common_proportion > 0.4: # Example threshold
                            categorical_char.append(f"- **{feature_name}**: Predominantly **{most_common_category}** ({most_common_proportion:.1%})")
                        else:
                            # If no single category is dominant, list top few or describe distribution
                            sorted_proportions = proportions_for_feature.sort_values(ascending=False)
                            top_categories = []
                            for cat_full, prop in sorted_proportions.items():
                                cat_name = cat_full.split(': ')[1] if ': ' in cat_full else cat_full
                                if prop > 0.1: # List categories with more than 10%
                                    top_categories.append(f"{cat_name} ({prop:.1%})")
                            if top_categories:
                                categorical_char.append(f"- **{feature_name}**: Diverse, with top categories: {', '.join(top_categories)}")


        # Generic persona/implications based on characteristics
        persona = "This cluster represents a diverse group with varied characteristics."
        if numeric_char or categorical_char:
            persona = "This cluster is characterized by:"
            if numeric_char:
                persona += "\n" + "\n".join(numeric_char)
            if categorical_char:
                persona += "\n" + "\n".join(categorical_char)
        
        structured_summaries[cluster_id_str] = {
            "cluster_heading": f"Cluster {cluster_id_int}: (N={cluster_size}, {cluster_percentage:.1f}% of total data)",
            "numeric_characteristics": numeric_char,
            "categorical_characteristics": categorical_char,
            "persona_implications": persona
        }
    
    # Add a summary for DBSCAN noise if it exists
    if chosen_algo == "DBSCAN" and -1 in unique_labels:
        noise_size = label_counts.get(-1, 0)
        noise_percentage = (noise_size / total_samples) * 100 if total_samples > 0 else 0
        structured_summaries["-1"] = {
            "cluster_heading": f"Noise (Unclustered) Samples: (N={noise_size}, {noise_percentage:.1f}% of total data)",
            "numeric_characteristics": [],
            "categorical_characteristics": [],
            "persona_implications": "These samples were identified as noise by DBSCAN and do not belong to any cluster."
        }

    return structured_summaries

# Function to create a Word report
def create_report(document, chosen_algo, algo_params, metrics, df_preview, pca_plot_bytes, profile_plot_bytes, cluster_means_numeric, cluster_cat_proportions, df_data_with_labels, labels, chosen_algo_for_summary):
    
    # Title
    document.add_heading("ML Analysis Report", level=0)
    document.add_paragraph(f"Report generated on: {pd.to_datetime('today').strftime('%Y-%m-%d %H:%M:%S')}")

    # 1. Analysis Overview
    document.add_heading("1. Analysis Overview", level=1)
    document.add_paragraph("This report details the unsupervised learning analysis performed using the Streamlit application. The goal is to group similar data points based on their attributes, enabling targeted strategies.")

    # 2. Clustering Parameters
    document.add_heading("2. Clustering Parameters", level=1)
    document.add_paragraph(f"Algorithm Used: {chosen_algo}")
    for param, value in algo_params.items():
        document.add_paragraph(f"- {param.replace('_', ' ').title()}: {value}")

    # 3. Model Performance Metrics
    document.add_heading("3. Model Performance Metrics", level=1)
    for metric, value in metrics.items():
        document.add_paragraph(f"{metric.replace('_', ' ').title()}: {value:.4f}")
    document.add_paragraph("These metrics evaluate the quality of the clusters. Higher Silhouette and Calinski-Harabasz scores indicate better-defined clusters. Lower Davies-Bouldin scores indicate better separation between clusters.")

    # 4. Data Preview
    document.add_heading("4. Data Preview (First 5 Rows)", level=1)
    if df_preview is not None and not df_preview.empty:
        table = document.add_table(rows=1, cols=len(df_preview.columns))
        table.style = 'Table Grid'
        # Add header row
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df_preview.columns):
            hdr_cells[i].text = col_name
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add data rows
        for index, row in df_preview.iterrows():
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row):
                row_cells[i].text = str(cell_value)
    else:
        document.add_paragraph("No data preview available.")
    document.add_paragraph("") # Add a blank line after the table

    # 5. Visualizations
    document.add_heading("5. Visualizations", level=1)
    if pca_plot_bytes:
        document.add_paragraph("PCA Plot of Clusters:")
        document.add_picture(io.BytesIO(pca_plot_bytes), width=Inches(6.0))
    else:
        document.add_paragraph("PCA Plot not available.")

    if profile_plot_bytes:
        document.add_paragraph("Average values of numeric features for each cluster:")
        document.add_picture(io.BytesIO(profile_plot_bytes), width=Inches(6.0))
    else:
        document.add_paragraph("Numeric Feature Profile Plot not available.")
    document.add_paragraph("") # Add a blank line

    # 6. Cluster Profiles (Tables)
    document.add_heading("6. Cluster Profiles (Tables)", level=1)
    document.add_paragraph("Below are the average values for numeric features and proportions for categorical features within each cluster, providing a quantitative profile of each segment.")

    # Numeric Feature Means Table
    document.add_heading("Numeric Feature Means by Cluster", level=2)
    if cluster_means_numeric is not None and not cluster_means_numeric.empty:
        table = document.add_table(rows=1, cols=len(cluster_means_numeric.columns) + 1)
        table.style = 'Table Grid'
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Cluster'
        for i, col_name in enumerate(cluster_means_numeric.columns):
            hdr_cells[i+1].text = col_name
            for paragraph in hdr_cells[i+1].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        # Data rows
        for cluster_id, row in cluster_means_numeric.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(cluster_id)
            for i, cell_value in enumerate(row):
                row_cells[i+1].text = format_number(cell_value)
    else:
        document.add_paragraph("No numeric feature means data available.")
    document.add_paragraph("")

    # Categorical Feature Proportions Table
    document.add_heading("Categorical Feature Proportions by Cluster", level=2)
    if cluster_cat_proportions is not None and not cluster_cat_proportions.empty:
        # Re-index to ensure correct order
        cluster_cat_proportions = cluster_cat_proportions.sort_index()

        table = document.add_table(rows=1, cols=len(cluster_cat_proportions.columns) + 1)
        table.style = 'Table Grid'
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Cluster'
        for i, col_name in enumerate(cluster_cat_proportions.columns):
            hdr_cells[i+1].text = col_name
            for paragraph in hdr_cells[i+1].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        # Data rows
        for cluster_id, row in cluster_cat_proportions.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(cluster_id)
            for i, cell_value in enumerate(row):
                row_cells[i+1].text = format_percentage(cell_value)
    else:
        document.add_paragraph("No categorical feature proportions data available.")
    document.add_paragraph("")

    # 7. Cluster Summaries and Characteristics (Textual)
    document.add_heading("7. Cluster Summaries and Characteristics", level=1)
    document.add_paragraph("Below is a detailed profile for each identified cluster, highlighting their key distinguishing features compared to the overall dataset. This can help you understand the unique characteristics of each segment and develop targeted strategies.")
    
    # This is the function that needs the df_data_with_labels argument updated.
    # Pass df_data_with_labels (which has original features and cluster labels)
    cluster_summaries = generate_cluster_summaries(
        cluster_means_numeric, 
        cluster_cat_proportions, 
        df_data_with_labels, # This is the crucial change
        labels,
        chosen_algo_for_summary
    )

    for cluster_id_str in sorted(cluster_summaries.keys(), key=lambda x: int(x) if x != '-1' else float('inf')):
        summary = cluster_summaries[cluster_id_str]
        document.add_heading(summary["cluster_heading"], level=3)
        
        # Add numeric characteristics
        if summary["numeric_characteristics"]:
            document.add_paragraph("Numeric Characteristics:")
            for char in summary["numeric_characteristics"]:
                document.add_paragraph(char, style='List Bullet')
        
        # Add categorical characteristics
        if summary["categorical_characteristics"]:
            document.add_paragraph("Categorical Characteristics:")
            for char in summary["categorical_characteristics"]:
                document.add_paragraph(char, style='List Bullet')
        
        # Add persona/implications
        if summary["persona_implications"]:
            if not summary["numeric_characteristics"] and not summary["categorical_characteristics"]:
                 # If no specific characteristics, just print the default message
                 document.add_paragraph(summary["persona_implications"])
            else:
                # If there are characteristics, add a general intro
                document.add_paragraph("Key Implications and Persona:")
                document.add_paragraph(summary["persona_implications"])

        document.add_paragraph("") # Blank line after each cluster summary

# --- Main Streamlit Application ---
def main_app():
    # Attempt to load logo for main app
    try:
        logo = Image.open("SsoLogo.jpg")
        st.sidebar.image(logo, width=150) # Adjust width as needed
    except FileNotFoundError:
        st.sidebar.warning("SsoLogo.jpg not found. Please ensure it's in the same directory as the script.")
    
    st.sidebar.title("Unsupervised Learning (ML)")
    st.sidebar.markdown("Perform customer segmentation using various unsupervised learning algorithms.")

    st.title("Customer Segmentation with Unsupervised Learning")

    df = None
    if not st.session_state.analysis_completed:
        with st.expander("Upload Data", expanded=True):
            uploaded_file = st.file_uploader(
                "Upload your customer data (.csv or .xlsx)", 
                type=["csv", "xlsx"], 
                key=st.session_state.file_uploader_key
            )
            if uploaded_file is not None:
                df = load_data(uploaded_file)
                if df is not None:
                    st.success("Data loaded successfully!")
                    st.subheader("Data Preview (First 5 Rows)")
                    st.dataframe(df.head())

                    # Suggest ID columns based on common patterns or unique values
                    id_cols = [col for col in df.columns if 'id' in col.lower() or df[col].nunique() == len(df)]
                    st.info(f"Detected potential ID columns: {', '.join(id_cols) if id_cols else 'None'}. Please ensure they are excluded from analysis.")

            else:
                st.info("Please upload a data file to begin.")
                return # Stop execution if no file uploaded

    if df is not None:
        with st.expander("Data Preprocessing & Feature Selection", expanded=not st.session_state.analysis_completed):
            all_columns = df.columns.tolist()
            
            # Allow exclusion of columns
            cols_to_exclude = st.multiselect(
                "Select columns to EXCLUDE from analysis (e.g., IDs, irrelevant text)",
                all_columns,
                default=[col for col in all_columns if 'id' in col.lower()]
            )
            
            available_cols_for_selection = [col for col in all_columns if col not in cols_to_exclude]

            numeric_cols = df[available_cols_for_selection].select_dtypes(include=np.number).columns.tolist()
            categorical_cols = df[available_cols_for_selection].select_dtypes(include='object').columns.tolist()

            st.subheader("Feature Selection for Clustering")
            selected_numeric = st.multiselect(
                "Select Numeric Features for Clustering:",
                numeric_cols,
                default=numeric_cols # Pre-select all numeric by default
            )
            selected_categorical = st.multiselect(
                "Select Categorical Features for Clustering (will be One-Hot Encoded):",
                categorical_cols,
                default=categorical_cols # Pre-select all categorical by default
            )

            if not selected_numeric and not selected_categorical:
                st.warning("Please select at least one numeric or categorical feature to proceed.")
                return

            st.subheader("Missing Data Handling")
            missing_data_strategy = st.selectbox(
                "Choose strategy for missing values in selected features:",
                ["drop_rows", "mean_imputation", "median_imputation", "mode_imputation"],
                help="Drop Rows: Removes rows with any missing values. Mean/Median Imputation: Fills numeric missing values with mean/median. Mode Imputation: Fills categorical missing values with mode."
            )

            st.subheader("Outlier Handling (Numeric Features Only)")
            outlier_handling_method = st.selectbox(
                "Choose strategy for outliers in numeric features:",
                ["none", "iqr_capping"],
                help="IQR Capping: Values outside 1.5*IQR bounds are capped to the bounds."
            )

            if st.button("Apply Preprocessing & Prepare for Clustering"):
                st.session_state.analysis_completed = False # Reset analysis completion status
                with st.spinner("Applying preprocessing..."):
                    try:
                        scaled_df, df_profile = preprocess_data(df, selected_numeric, selected_categorical, missing_data_strategy, outlier_handling_method)
                        st.session_state.scaled_df = scaled_df
                        st.session_state.df_profile = df_profile # Original numeric/categorical values for profiling
                        st.session_state.selected_numeric = selected_numeric # Store for later plotting
                        st.session_state.selected_categorical = selected_categorical # Store for later profiling
                        st.success("Data preprocessed and ready for clustering!")
                    except Exception as e:
                        st.error(f"Error during preprocessing: {e}")
                        st.session_state.scaled_df = None
                        st.session_state.df_profile = None

        if "scaled_df" in st.session_state and st.session_state.scaled_df is not None:
            scaled_df = st.session_state.scaled_df
            df_profile = st.session_state.df_profile
            selected_numeric = st.session_state.selected_numeric
            selected_categorical = st.session_state.selected_categorical

            with st.expander("Clustering Evaluation (Find Optimal K)", expanded=False):
                st.subheader("Evaluate Algorithms & Number of Clusters (K)")
                k_range = st.slider("Select K Range for Evaluation:", min_value=2, max_value=10, value=(2, 5))

                run_eval = st.button("Run Evaluation for KMeans, GMM, Agglomerative")
                if run_eval:
                    k_values = range(k_range[0], k_range[1] + 1)
                    results = []
                    for k in k_values:
                        st.info(f"Evaluating K={k}...")
                        for algo_name, algo_model in {
                            "KMeans": KMeans(n_clusters=k, random_state=42, n_init='auto'),
                            "GaussianMixture": GaussianMixture(n_components=k, random_state=42),
                            "AgglomerativeClustering": AgglomerativeClustering(n_clusters=k)
                        }.items():
                            try:
                                labels = algo_model.fit_predict(scaled_df)
                                if len(np.unique(labels)) > 1: # Metrics require > 1 cluster
                                    sil = silhouette_score(scaled_df, labels)
                                    db = davies_bouldin_score(scaled_df, labels)
                                    ch = calinski_harabasz_score(scaled_df, labels)
                                    results.append({"K": k, "Algorithm": algo_name, "Silhouette": sil, "Davies-Bouldin": db, "Calinski-Harabasz": ch})
                                else:
                                    results.append({"K": k, "Algorithm": algo_name, "Silhouette": np.nan, "Davies-Bouldin": np.nan, "Calinski-Harabasz": np.nan})
                            except Exception as e:
                                st.warning(f"Could not evaluate {algo_name} for K={k}: {e}")
                                results.append({"K": k, "Algorithm": algo_name, "Silhouette": np.nan, "Davies-Bouldin": np.nan, "Calinski-Harabasz": np.nan})

                    if results:
                        results_df = pd.DataFrame(results)
                        st.dataframe(results_df.set_index(['K', 'Algorithm']))

                        # Plotting metrics
                        fig, axes = plt.subplots(1, 3, figsize=(18, 5))
                        sns.lineplot(data=results_df, x='K', y='Silhouette', hue='Algorithm', marker='o', ax=axes[0])
                        axes[0].set_title('Silhouette Score')
                        axes[0].set_ylabel('Score (Higher is Better)')

                        sns.lineplot(data=results_df, x='K', y='Davies-Bouldin', hue='Algorithm', marker='o', ax=axes[1])
                        axes[1].set_title('Davies-Bouldin Index')
                        axes[1].set_ylabel('Score (Lower is Better)')

                        sns.lineplot(data=results_df, x='K', y='Calinski-Harabasz', hue='Algorithm', marker='o', ax=axes[2])
                        axes[2].set_title('Calinski-Harabasz Index')
                        axes[2].set_ylabel('Score (Higher is Better)')

                        plt.tight_layout()
                        st.pyplot(fig)
                        plt.close(fig)

                        # Elbow Method for KMeans
                        if "KMeans" in results_df["Algorithm"].unique():
                            inertia = []
                            for k in k_values:
                                kmeans_model = KMeans(n_clusters=k, random_state=42, n_init='auto')
                                kmeans_model.fit(scaled_df)
                                inertia.append(kmeans_model.inertia_)
                            
                            fig_elbow, ax_elbow = plt.subplots(figsize=(8, 5))
                            ax_elbow.plot(k_values, inertia, marker='o')
                            ax_elbow.set_title('Elbow Method for KMeans (Inertia)')
                            ax_elbow.set_xlabel('Number of Clusters (K)')
                            ax_elbow.set_ylabel('Inertia')
                            st.pyplot(fig_elbow)
                            plt.close(fig_elbow)
                        
                        # AIC/BIC for GMM
                        if "GaussianMixture" in results_df["Algorithm"].unique():
                            aic_scores = []
                            bic_scores = []
                            for k in k_values:
                                gmm_model = GaussianMixture(n_components=k, random_state=42)
                                gmm_model.fit(scaled_df)
                                aic_scores.append(gmm_model.aic(scaled_df))
                                bic_scores.append(gmm_model.bic(scaled_df))
                            
                            fig_gmm_criteria, ax_gmm_criteria = plt.subplots(figsize=(8, 5))
                            ax_gmm_criteria.plot(k_values, aic_scores, marker='o', label='AIC')
                            ax_gmm_criteria.plot(k_values, bic_scores, marker='o', label='BIC')
                            ax_gmm_criteria.set_title('AIC/BIC for Gaussian Mixture Models')
                            ax_gmm_criteria.set_xlabel('Number of Components (K)')
                            ax_gmm_criteria.set_ylabel('Score (Lower is Better)')
                            ax_gmm_criteria.legend()
                            st.pyplot(fig_gmm_criteria)
                            plt.close(fig_gmm_criteria)

                        # Basic recommendation
                        best_silhouette = results_df.loc[results_df['Silhouette'].idxmax()]
                        best_db = results_df.loc[results_df['Davies-Bouldin'].idxmin()]
                        best_ch = results_df.loc[results_df['Calinski-Harabasz'].idxmax()]

                        st.markdown("---")
                        st.subheader("Recommendation based on Metrics:")
                        st.info(f"**Highest Silhouette Score**: {best_silhouette['Algorithm']} with K={best_silhouette['K']:.0f} (Score: {best_silhouette['Silhouette']:.4f})")
                        st.info(f"**Lowest Davies-Bouldin Index**: {best_db['Algorithm']} with K={best_db['K']:.0f} (Score: {best_db['Davies-Bouldin']:.4f})")
                        st.info(f"**Highest Calinski-Harabasz Index**: {best_ch['Algorithm']} with K={best_ch['K']:.0f} (Score: {best_ch['Calinski-Harabasz']:.4f})")
                        st.markdown("Consider these metrics along with domain knowledge to choose the optimal K and algorithm.")
                    else:
                        st.warning("No evaluation results. Please check your data and selections.")
            
            with st.expander("Run Final Clustering & Get Results", expanded=True):
                st.subheader("Select Algorithm and Parameters for Final Clustering")
                chosen_algo = st.selectbox(
                    "Select Clustering Algorithm:",
                    ["KMeans", "Gaussian Mixture Model", "Agglomerative Clustering", "DBSCAN"]
                )

                n_clusters = None
                eps = None
                min_samples = None

                if chosen_algo in ["KMeans", "Gaussian Mixture Model", "Agglomerative Clustering"]:
                    n_clusters = st.slider("Number of Clusters (K):", min_value=1 if chosen_algo == "Gaussian Mixture Model" else 2, max_value=15, value=5)
                elif chosen_algo == "DBSCAN":
                    eps = st.slider("DBSCAN - Epsilon (eps): Max distance between two samples for one to be considered as in the neighborhood of the other.",
                                    min_value=0.1, max_value=5.0, value=0.5, step=0.1)
                    min_samples = st.slider("DBSCAN - Min Samples: Number of samples in a neighborhood for a point to be considered as a core point.",
                                            min_value=1, max_value=20, value=5)

                if st.button("🚀 Run Clustering & Generate Report"):
                    labels = run_clustering(scaled_df, n_clusters, chosen_algo, eps, min_samples)
                    
                    if labels is not None:
                        unique_clusters = np.unique(labels)
                        num_clusters_formed = len(unique_clusters) if -1 not in unique_clusters else len(unique_clusters) - 1
                        
                        if num_clusters_formed >= 1: # Ensure at least one actual cluster (excluding noise)
                            silhouette, davies_bouldin, calinski_harabasz = evaluate_clustering(scaled_df, labels)

                            st.session_state.analysis_completed = True
                            st.session_state.labels = labels
                            st.session_state.silhouette = silhouette
                            st.session_state.davies_bouldin = davies_bouldin
                            st.session_state.calinski_harabasz = calinski_harabasz
                            st.session_state.chosen_algo = chosen_algo
                            st.session_state.algo_params = {'n_clusters': n_clusters, 'eps': eps, 'min_samples': min_samples}
                            st.session_state.n_clusters = n_clusters # Store n_clusters for report even if not used by DBSCAN
                            
                            st.subheader("Clustering Results Summary")
                            st.write(f"**Algorithm Used**: {chosen_algo}")
                            if n_clusters: st.write(f"**Number of Clusters (K)**: {n_clusters}")
                            if eps: st.write(f"**Epsilon (eps)**: {eps}")
                            if min_samples: st.write(f"**Min Samples** : {min_samples}")

                            st.write(f"**Silhouette Score**: {silhouette:.4f}")
                            st.write(f"**Davies-Bouldin Index**: {davies_bouldin:.4f}")
                            st.write(f"**Calinski-Harabasz Index**: {calinski_harabasz:.4f}")

                            st.subheader("Cluster Distribution")
                            cluster_counts = pd.Series(labels).value_counts().sort_index()
                            st.dataframe(cluster_counts.rename("Count").reset_index().rename(columns={'index': 'Cluster'}))

                            # Prepare data for plots and report
                            df_profile_with_labels = df_profile.copy()
                            df_profile_with_labels['Cluster'] = labels # Keep labels as integers here

                            fig_pca, fig_profile_numeric, cluster_means_numeric, cluster_cat_proportions = generate_plots(
                                df_profile_with_labels, labels, selected_numeric
                            )

                            # Save plots to bytes for report
                            pca_plot_bytes = io.BytesIO()
                            if fig_pca:
                                fig_pca.savefig(pca_plot_bytes, format='png', bbox_inches='tight')
                                pca_plot_bytes.seek(0)
                                st.pyplot(fig_pca)
                            
                            profile_plot_bytes = io.BytesIO()
                            if fig_profile_numeric:
                                fig_profile_numeric.savefig(profile_plot_bytes, format='png', bbox_inches='tight')
                                profile_plot_bytes.seek(0)
                                st.pyplot(fig_profile_numeric)

                            # Generate Word report
                            document = Document()
                            create_report(
                                document,
                                chosen_algo,
                                {'n_clusters': n_clusters, 'eps': eps, 'min_samples': min_samples} if chosen_algo != 'DBSCAN' else {'eps': eps, 'min_samples': min_samples},
                                {'silhouette': silhouette, 'davies_bouldin': davies_bouldin, 'calinski_harabasz': calinski_harabasz},
                                df_profile_with_labels.head(5), # Data preview
                                pca_plot_bytes.getvalue() if pca_plot_bytes else None,
                                profile_plot_bytes.getvalue() if profile_plot_bytes else None,
                                cluster_means_numeric,
                                cluster_cat_proportions,
                                df_profile_with_labels, # Pass df_profile_with_labels for accurate summaries
                                labels,
                                chosen_algo # Pass chosen_algo for DBSCAN noise handling
                            )
                            
                            report_bytes = io.BytesIO()
                            document.save(report_bytes)
                            report_bytes.seek(0)

                            st.success("Analysis Complete! Download your reports below.")
                            st.download_button(
                                label="📥 Download Comprehensive Report (.docx)",
                                data=report_bytes,
                                file_name=f"ML_Analysis_Report_{pd.to_datetime('today').strftime('%Y%m%d_%H%M%S')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

                            csv_buffer = io.StringIO()
                            df_profile_with_labels.to_csv(csv_buffer, index=False)
                            st.download_button(
                                label="📥 Download Clustered Data (CSV)",
                                data=csv_buffer.getvalue(),
                                file_name=f"Clustered_Customer_Data_{pd.to_datetime('today').strftime('%Y%m%d_%H%M%S')}.csv",
                                mime="text/csv",
                            )
                        else:
                            st.error("Clustering did not form more than one cluster or failed. Please review your data and selected parameters.")
                            st.session_state.analysis_completed = False

    # Reset Buttons and "What's Next" section
    if st.session_state.analysis_completed:
        st.header("🎯 Analysis Complete")
        st.markdown("You can either re-run clustering with different parameters on the current dataset, or clear everything to start fresh with new data.")
        col_reset1, col_reset2 = st.columns(2)
        with col_reset1:
            if st.button("🔄 Rerun Clustering with New Parameters"):
                st.session_state.analysis_completed = False
                st.rerun()
        with col_reset2:
            if st.button("🗑️ Clear All Data & Start Fresh"):
                current_file_uploader_key = st.session_state.get('file_uploader_key', 0)
                st.session_state.clear()
                st.session_state.file_uploader_key = current_file_uploader_key + 1
                st.rerun()
    else:
        if df is None:
            st.info("Please upload a data file (.csv or .xlsx) at the top of the page to begin the analysis.")

# --- Run the appropriate page based on authentication status ---
if not st.session_state.authenticated:
    login_page()
else:
    main_app()
